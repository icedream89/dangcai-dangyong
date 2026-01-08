# -*- coding: utf-8 -*-
"""
生成项目规划方案Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell):
    """设置单元格边框"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)

    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, font_size=11):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True
    return para

def add_bullet_point(doc, text, level=0):
    """添加列表项"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return para

def create_document():
    """创建Word文档"""
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(11)

    # 标题
    title = doc.add_heading('"当才当用"政府服务平台', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_heading('开发实施计划', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 项目概述
    add_heading(doc, '一、项目概述', level=1)
    add_paragraph(doc, '项目名称：当才当用', bold=True)
    add_paragraph(doc, '客户：当阳市科技经济信息商务局', bold=True)
    add_paragraph(doc, '项目目标：搭建政府-企业-用户三方服务平台', bold=True)

    add_paragraph(doc, '\n技术栈：', bold=True)
    tech_stack = [
        '后台管理前端: Vue 3 + Element Plus',
        '后端: Java (Spring Boot)',
        '数据库: MySQL 8.0',
        '小程序: uni-app',
        '缓存: Redis 7.x',
        '文件存储: MinIO'
    ]
    for tech in tech_stack:
        add_bullet_point(doc, tech)

    # 系统架构
    add_heading(doc, '二、整体系统架构', level=1)

    add_paragraph(doc, '2.1 架构分层', bold=True)
    arch_desc = [
        '客户端层: Web后台管理 + uni-app微信小程序',
        '应用服务层: Spring Boot + Spring Security',
        '数据层: MySQL + Redis + MinIO'
    ]
    for desc in arch_desc:
        add_bullet_point(doc, desc)

    add_paragraph(doc, '\n2.2 后端模块划分', bold=True)
    modules = [
        'server-common - 通用模块（工具类、异常处理、基础配置）',
        'server-system - 系统模块（用户、角色、菜单、权限）',
        'server-enterprise - 企业模块（企业管理、员工管理）',
        'server-business - 业务模块（分类、产品、政策、课堂、工单、需求）',
        'server-file - 文件服务模块',
        'server-message - 消息服务模块',
        'server-admin - 后台管理API启动模块',
        'server-miniapp - 小程序API启动模块'
    ]
    for module in modules:
        add_bullet_point(doc, module)

    # 数据库设计
    add_heading(doc, '三、数据库核心表设计', level=1)

    add_paragraph(doc, '系统共设计22张核心表：', bold=True)

    add_paragraph(doc, '\n3.1 系统核心表（6张）', bold=True)
    sys_tables = [
        'sys_user - 用户表（统一用户认证）',
        'sys_role - 角色表',
        'sys_menu - 菜单权限表',
        'sys_user_role - 用户角色关联',
        'sys_role_menu - 角色菜单关联',
        'sys_config - 系统配置'
    ]
    for table in sys_tables:
        add_bullet_point(doc, table)

    add_paragraph(doc, '\n3.2 企业核心表（3张）', bold=True)
    ent_tables = [
        'ent_enterprise - 企业表',
        'ent_employee - 企业员工表',
        'ent_enterprise_log - 企业变更记录'
    ]
    for table in ent_tables:
        add_bullet_point(doc, table)

    add_paragraph(doc, '\n3.3 业务核心表（9张）', bold=True)
    biz_tables = [
        'biz_category - 分类表（树形结构）',
        'biz_product - 产品表',
        'biz_policy - 政策表',
        'biz_course - 企业课堂表',
        'biz_help_ticket - 求助工单表',
        'biz_help_flow - 工单流转记录表',
        'biz_purchase_requirement - 采购需求表',
        'biz_purchase_reply - 采购需求回复表',
        'biz_message_subscribe - 消息订阅表'
    ]
    for table in biz_tables:
        add_bullet_point(doc, table)

    add_paragraph(doc, '\n3.4 日志表（3张）', bold=True)
    log_tables = [
        'sys_oper_log - 操作日志',
        'sys_login_log - 登录日志',
        'sys_file - 文件记录'
    ]
    for table in log_tables:
        add_bullet_point(doc, table)

    # 开发阶段规划
    add_heading(doc, '四、开发阶段规划', level=1)

    add_paragraph(doc, '项目分三个阶段开发，预计周期12-17周', bold=True)

    add_paragraph(doc, '\n阶段一: MVP阶段（4-6周）', bold=True)
    mvp_tasks = [
        '第1周: 基础架构搭建 - 初始化前后端项目，创建数据库，配置中间件',
        '第2周: 系统管理模块 - 用户、角色、菜单管理，JWT认证',
        '第3周: 企业管理模块 - 企业信息、员工管理，变更记录追踪',
        '第4周: 业务模块基础功能 - 分类、政策、课堂、产品管理',
        '第5周: 小程序核心功能 - 微信登录、首页、政策、课堂、企业页面'
    ]
    for task in mvp_tasks:
        add_bullet_point(doc, task)

    add_paragraph(doc, '\n阶段二: 完整功能阶段（6-8周）', bold=True)
    full_tasks = [
        '第6-7周: 业务模块高级功能 - 求助工单系统、采购需求、消息订阅',
        '第8-9周: 小程序完整功能 - 我要求助、我要采购、个人中心',
        '第10-11周: 后台管理完善功能 - 系统配置、文件管理、批量导入导出、数据统计',
        '第12-13周: 性能优化与测试 - 接口优化、前端优化、单元测试、集成测试、压力测试'
    ]
    for task in full_tasks:
        add_bullet_point(doc, task)

    add_paragraph(doc, '\n阶段三: 部署上线阶段（2-4周）', bold=True)
    deploy_tasks = [
        '第14周: 部署准备 - 服务器环境搭建、Nginx配置、SSL证书、域名配置',
        '第15周: 小程序上线 - 代码审核、域名配置、小程序发布',
        '第16-17周: 试运行与优化 - 数据迁移、用户培训、问题修复、性能调优'
    ]
    for task in deploy_tasks:
        add_bullet_point(doc, task)

    # 关键技术
    add_heading(doc, '五、关键技术实现要点', level=1)

    tech_points = [
        '1. 认证授权方案（JWT + RBAC）- 使用JWT进行用户认证，基于角色的访问控制，数据权限控制',
        '2. 树形结构分类管理 - 使用parent_id + path字段实现无限层级，递归构建树形结构',
        '3. 文件上传方案（MinIO）- 支持多格式文件，分业务类型存储',
        '4. 工单流转系统 - 状态机设计，工单流转记录追踪，时限管理',
        '5. 微信订阅消息 - 用户订阅管理，模板消息推送',
        '6. 批量导入导出（EasyExcel）- Excel模板下载，数据校验，分批处理'
    ]
    for point in tech_points:
        add_bullet_point(doc, point)

    # 开发优先级
    add_heading(doc, '六、开发优先级', level=1)

    add_paragraph(doc, '\nP0 - 核心功能（必须完成）', bold=True)
    p0_tasks = [
        '用户认证与授权',
        '企业管理',
        '分类管理',
        '政策管理',
        '小程序微信登录',
        '小程序基础页面'
    ]
    for task in p0_tasks:
        add_bullet_point(doc, task, level=1)

    add_paragraph(doc, '\nP1 - 重要功能（尽快完成）', bold=True)
    p1_tasks = [
        '企业课堂',
        '产品管理',
        '求助工单系统',
        '采购需求管理',
        '文件上传',
        '批量导入导出'
    ]
    for task in p1_tasks:
        add_bullet_point(doc, task, level=1)

    add_paragraph(doc, '\nP2 - 增强功能（逐步完成）', bold=True)
    p2_tasks = [
        '微信订阅消息',
        '数据统计分析',
        '操作日志',
        '系统配置'
    ]
    for task in p2_tasks:
        add_bullet_point(doc, task, level=1)

    add_paragraph(doc, '\nP3 - 可选功能（根据时间）', bold=True)
    p3_tasks = [
        '数据大屏',
        '消息通知中心',
        '性能监控'
    ]
    for task in p3_tasks:
        add_bullet_point(doc, task, level=1)

    # 风险与注意事项
    add_heading(doc, '七、风险与注意事项', level=1)

    risks = [
        '技术风险: MinIO学习成本、树形结构性能、微信订阅消息限制',
        '业务风险: 工单流转逻辑复杂、数据权限控制、批量导入导出数据校验',
        '性能风险: 大文件上传、Excel导入、数据库查询优化',
        '开发建议: 先完成后优化、代码复用、注释清晰、定期代码审查'
    ]
    for risk in risks:
        add_bullet_point(doc, risk)

    # 项目状态
    add_heading(doc, '八、当前项目状态', level=1)

    add_paragraph(doc, '\n已完成工作：', bold=True)
    completed = [
        '✅ 完整系统架构设计（22张数据库表）',
        '✅ 技术栈选型（Vue3 + Spring Boot + MySQL + uni-app）',
        '✅ 管理后台Demo（13个功能页面，可运行演示）',
        '✅ 小程序Demo（9个功能页面，代码已完成）'
    ]
    for item in completed:
        add_bullet_point(doc, item)

    add_paragraph(doc, '\n待完成工作：', bold=True)
    pending = [
        '⬜ 后端开发（Spring Boot + MySQL）',
        '⬜ 前后端联调',
        '⬜ 小程序显示问题修复',
        '⬜ 性能优化',
        '⬜ 测试和上线'
    ]
    for item in pending:
        add_bullet_point(doc, item)

    # 尾页信息
    doc.add_page_break()
    add_paragraph(doc, '\n\n', bold=False)
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f'计划制定时间: 2026-01-06\n')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)

    info_para = doc.add_paragraph()
    info_run = info_para.add_run('预计开发周期: 12-17周\n')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)

    info_para = doc.add_paragraph()
    info_run = info_para.add_run('开发优先级: 按P0 → P1 → P2 → P3顺序开发\n')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)

    return doc

if __name__ == '__main__':
    try:
        doc = create_document()
        output_path = r'C:\Users\icedr\Desktop\test\dang-cai-dang-yong\项目规划方案.docx'
        doc.save(output_path)
        print('[SUCCESS] Word document generated successfully!')
        print(f'[FILE] Location: {output_path}')
    except ImportError:
        print('[ERROR] python-docx library not installed')
        print('Please run: pip install python-docx')
    except Exception as e:
        print(f'[ERROR] Generation failed: {str(e)}')
